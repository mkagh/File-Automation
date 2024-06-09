import PySimpleGUI as sg
import camelot as cm
import pandas as pd
import pdfkit as pdf
import os
from docx import Document
def PDF(MainMenuWindow):
 def createGUI():
    layout=[
      
        [ sg.Input(key="-FILE-", enable_events=True),sg.FileBrowse("browse pdf",key="-browsePDF-" ,target="-FILE-",button_color="red")],
        [sg.Input("Choose table",key="-NUMBERofTABLE-")],
        [sg.Input("From Column", key="-FCOL-")],
         [sg.Input("To Column",key="-TCOL-")], 
        [sg.Input("From Row", key="-FROW-")],
         [sg.Input("To Row",key="-TROW-")], 
         [sg.Button("Process",key="-PROCESS-",disabled=True,button_color="red")],
        [sg.Button("To CSV",key="-TOCSV-",disabled=True,button_color="red"),
         sg.Button("To Word",key="-TOWORD-",disabled=True,button_color="red")],
         [sg.VPush()],
           [sg.Button("MainMenu",key="-MAINMENU-",size=30,button_color="red")]
    ]
    return sg.Window("File Automation",layout,size=(600,600),element_justification="center")

 window=createGUI()
 while True:
    event,values=window.read()
       
    if event== sg.WIN_CLOSED:
        break
    if event== "-FILE-":
       window["-PROCESS-"].update(disabled=False)
    if event== "-PROCESS-":
      _, file_extension = os.path.splitext(values["-browsePDF-"])
      if file_extension.lower() == '.pdf':
        wanted_file=cm.read_pdf(f"{values["-browsePDF-"]}",flavor="lattice",pages="1,2")
        if len(wanted_file)>0:
          try:
           if int(values['-NUMBERofTABLE-'])>len(wanted_file):
              table_index=0
           else:
              table_index = int(values['-NUMBERofTABLE-'])
          except ValueError:
            table_index = 0
            df=wanted_file[table_index].df
            def locationsStart(Value):
                try:
                    return int(values[Value])
                except ValueError:
                    if Value=="-FROW-" or Value=="-FCOL-":
                      return  0
                    elif Value=="-TROW-":
                      return  df.shape[0]
                    elif Value=="-TCOL-":
                      return  df.shape[1]
            proccesedDF=df.loc[locationsStart("-FROW-"):locationsStart("-TROW-"),locationsStart("-FCOL-"):locationsStart("-TCOL-")]
            print(proccesedDF)
            window["-TOCSV-"].update(disabled=False) 
            window["-TOWORD-"].update(disabled=False)
        else:
          window["-FILE-"].update("nema tabela u ovom fajlu")
      else:
         window["-FILE-"].update("This is not pdf file")
    if event== "-TOCSV-":
       proccesedDF.to_csv("table.csv")
    if event== "-TOWORD-":
        doc = Document()
        t=doc.add_table(proccesedDF.shape[0]+1,proccesedDF.shape[1])
        
        for j in range(proccesedDF.shape[-1]):
           t.cell(0,j).text = str(proccesedDF.columns[j])

        for i in range(proccesedDF.shape[0]):
            for j in range(proccesedDF.shape[-1]):
                t.cell(i+1,j).text = str(proccesedDF.values[i,j])   
         
        doc.save("test.docx")       
    if event== "-MAINMENU-":
        window.close()
        MainMenuWindow.un_hide()

def WORD(main_menu_window):
   def createWordGui():
      layout=[
              [sg.Input(key="-FILE-", enable_events=True),sg.FileBrowse("Browse Word",key="-browseWord-")],
                [sg.Button("process",key="-PROCESS-",disabled=True)],
                [
                 [sg.Button("tocsv",key="-TOCSV-",disabled=True),
                 sg.Button("HTML",key="-TOHTML-",disabled=True)]],
                 [sg.VPush()],
                 [sg.Button("MainMenu",key="-BACK-")]]
      return sg.Window("word",layout,size=(500,600),element_justification="center")
   window=createWordGui()
   while True:
      event,values=window.read()
      if event==sg.WIN_CLOSED:
        break
      if event=="-FILE-":
          window["-PROCESS-"].update(disabled=False)
      if event=="-PROCESS-":
        _, file_extension = os.path.splitext(values["-browseWord-"])
        if file_extension.lower() == '.docx':       
          doc=Document(values["-browseWord-"])
          if len(doc.tables)>0:
              table = doc.tables[0]
              start_row = 2
              data = [[cell.text for cell in row.cells] for i, row in enumerate(table.rows) 
              if i >= start_row]
              df = pd.DataFrame(data[1:], columns=data[0])
              window["-TOCSV-"].update(disabled=False) 
              window["-TOHTML-"].update(disabled=False)
          else:
           window["-FILE-"].update("There is no table in this file")
        else:
          window["-FILE-"].update("This is not word file")
      if event=="-TOCSV-":
        df.to_csv("tocsv.csv")
      if event=="-TOHTML-":
        df.to_html('izvestaj.html')
      if event=="-BACK-":
        window.close()
        main_menu_window.un_hide()
   
   window.close()   

def createMainMenu():
    sg.theme('Reddit')
    layout=[[sg.Push(),sg.Text("Chose what to convert:",text_color="black",
            font=(20)),sg.Push()],
            [sg.VPush()],
            [sg.Push(),sg.Button("PDF",key="-PDF-",size=10,button_color="red"),
             sg.Button("WORD",key="-WORD-",size=10,button_color="blue"),sg.Push()],
            [sg.VPush()]],
    return sg.Window("Main Menu",layout,size=(300,300))

window=createMainMenu()

while True:
    event,values=window.read()
    if event==sg.WIN_CLOSED:
        break
    if event=="-PDF-":
        window.hide()
        PDF(window)
    if event=="-WORD-":
        window.hide()
        WORD(window)
window.close()    
    
    